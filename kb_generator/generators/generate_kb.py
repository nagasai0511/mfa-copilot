import json
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm

from groq_client import generate

# ---------------------------------------------------------
# Logging
# ---------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s"
)

# ---------------------------------------------------------
# Directories
# ---------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent.parent

DATA_DIR = BASE_DIR / "data"
PROMPT_DIR = BASE_DIR / "prompts"
OUTPUT_DIR = BASE_DIR / "output"

OUTPUT_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------
# Load Prompt
# ---------------------------------------------------------

PROMPT_FILE = PROMPT_DIR / "knowledge_article.txt"

if not PROMPT_FILE.exists():
    raise FileNotFoundError(
        f"Prompt file not found: {PROMPT_FILE}"
    )

PROMPT_TEMPLATE = PROMPT_FILE.read_text(
    encoding="utf-8"
)

# ---------------------------------------------------------
# Load Incidents
# ---------------------------------------------------------

INCIDENT_FILE = DATA_DIR / "incidents.json"

if not INCIDENT_FILE.exists():
    raise FileNotFoundError(
        f"Incident dataset not found: {INCIDENT_FILE}"
    )

with open(INCIDENT_FILE, "r", encoding="utf-8") as f:
    incidents = json.load(f)

logging.info(f"Loaded {len(incidents)} incidents.")

# ---------------------------------------------------------
# Generate Documents
# ---------------------------------------------------------

for incident in tqdm(incidents):

    try:

        filename = (
            OUTPUT_DIR /
            f"{incident['id']}_{incident['title'].replace(' ','_')}.docx"
        )

        if filename.exists():

            logging.info(
                f"Skipping existing document : {filename.name}"
            )

            continue

        prompt = PROMPT_TEMPLATE.format(
            incident=json.dumps(
                incident,
                indent=4
            )
        )

        logging.info(
            f"Generating document for {incident['id']}"
        )

        article = generate(prompt)

        # ---------------------------------------------

        doc = Document()

        title = doc.add_heading(
            incident["title"],
            level=0
        )

        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        meta = doc.add_paragraph()

        meta.style.font.size = Pt(11)

        meta.add_run(
            f"Incident ID : {incident['id']}\n"
        )

        meta.add_run(
            f"Technology : {incident['technology']}\n"
        )

        meta.add_run(
            f"Severity : {incident['severity']}\n"
        )

        if "environment" in incident:
            meta.add_run(
                f"Environment : {incident['environment']}\n"
            )

        if "application" in incident:
            meta.add_run(
                f"Application : {incident['application']}\n"
            )

        doc.add_heading(
            "Knowledge Article",
            level=1
        )

        doc.add_paragraph(article)

        doc.save(filename)

        logging.info(
            f"Created : {filename.name}"
        )

    except Exception as e:

        logging.exception(
            f"Failed to generate {incident['id']} : {e}"
        )

logging.info("Knowledge Base Generation Completed.")
